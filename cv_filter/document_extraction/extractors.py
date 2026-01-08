"""Text extraction logic for different file formats."""

import logging
from typing import Dict, Any

import fitz  # PyMuPDF
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


class PDFExtractor:
    """Extracts text from PDF files using PyMuPDF."""

    @staticmethod
    def extract(file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary containing:
                - 'text': Extracted text content
                - 'metadata': File metadata
                - 'success': Boolean indicating extraction success
                - 'error': Error message if extraction failed
                - 'method': Extraction method used ('pymupdf')
        """
        try:
            text_parts = []
            metadata = {}

            with fitz.open(file_path) as doc:
                # Extract metadata
                metadata = {
                    "author": doc.metadata.get("author", ""),
                    "creator": doc.metadata.get("creator", ""),
                    "producer": doc.metadata.get("producer", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "title": doc.metadata.get("title", ""),
                    "page_count": doc.page_count,
                    "format": doc.metadata.get("format", ""),
                }

                # Extract text from all pages
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    page_text = page.get_text("text")

                    if page_text and page_text.strip():
                        text_parts.append(page_text)

            full_text = "\n".join(text_parts).strip()

            if full_text:
                logger.info(
                    f"Successfully extracted text using PyMuPDF from {file_path} "
                    f"({len(full_text)} chars)"
                )
                return {
                    "text": full_text,
                    "metadata": metadata,
                    "success": True,
                    "error": None,
                    "method": "pymupdf",
                }
            else:
                logger.warning(f"No text content extracted from PDF: {file_path}")
                return {
                    "text": "",
                    "metadata": metadata,
                    "success": False,
                    "error": "No text content extracted",
                    "method": "pymupdf",
                }

        except Exception as e:
            logger.error(f"PyMuPDF extraction error for {file_path}: {str(e)}")
            return {
                "text": "",
                "metadata": {},
                "success": False,
                "error": f"PyMuPDF error: {str(e)}",
                "method": "pymupdf",
            }


class DOCXExtractor:
    """Extracts text from DOCX files."""

    @staticmethod
    def extract(file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary containing:
                - 'text': Extracted text content
                - 'metadata': File metadata
                - 'success': Boolean indicating extraction success
                - 'error': Error message if extraction failed
                - 'method': Extraction method used ('python-docx')
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            # Extract headers and footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

            full_text = "\n".join(text_parts).strip()

            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                "author": core_properties.author if core_properties.author else "",
                "title": core_properties.title if core_properties.title else "",
                "subject": core_properties.subject if core_properties.subject else "",
                "created": (
                    str(core_properties.created) if core_properties.created else ""
                ),
                "modified": (
                    str(core_properties.modified) if core_properties.modified else ""
                ),
            }

            if full_text:
                logger.info(
                    f"Successfully extracted text from DOCX: {file_path} "
                    f"({len(full_text)} chars)"
                )
                return {
                    "text": full_text,
                    "metadata": metadata,
                    "success": True,
                    "error": None,
                    "method": "python-docx",
                }
            else:
                logger.warning(f"No text content extracted from DOCX: {file_path}")
                return {
                    "text": "",
                    "metadata": metadata,
                    "success": False,
                    "error": "No text content extracted",
                    "method": "python-docx",
                }

        except PackageNotFoundError:
            logger.error(f"Invalid DOCX file or file not found: {file_path}")
            return {
                "text": "",
                "metadata": {},
                "success": False,
                "error": "Invalid DOCX format or file not found",
                "method": "python-docx",
            }
        except Exception as e:
            logger.error(f"DOCX extraction error for {file_path}: {str(e)}")
            return {
                "text": "",
                "metadata": {},
                "success": False,
                "error": f"DOCX error: {str(e)}",
                "method": "python-docx",
            }
